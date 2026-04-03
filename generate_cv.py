"""
generate_cv.py
Generates Ikhlas Retbi's CV in both English and French as .docx files.
Run: python generate_cv.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_font(run, name='Calibri', size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, name='Calibri', size=10, bold=True, color=(0, 0, 0))
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_exp_item(doc, role, company, date, bullets):
    # Role + date on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(role)
    set_font(r1, size=10, bold=True)
    r2 = p.add_run(f'  —  {date}')
    set_font(r2, size=9, color=(100, 100, 100))

    # Company
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(company)
    set_font(r, size=9, italic=True, color=(80, 80, 80))

    # Bullets
    for b in bullets:
        p3 = doc.add_paragraph(style='List Bullet')
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(1)
        p3.paragraph_format.left_indent = Inches(0.2)
        r = p3.add_run(b)
        set_font(r, size=9, color=(50, 50, 50))


def add_skill_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{label}: ')
    set_font(r1, size=9, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=9, color=(60, 60, 60))


def add_project(doc, name, link, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(name)
    set_font(r1, size=10, bold=True)
    r2 = p.add_run(f'  ·  {link}')
    set_font(r2, size=8, color=(100, 100, 100))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(desc)
    set_font(r, size=9, color=(70, 70, 70))


def add_edu(doc, degree, school, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(degree)
    set_font(r1, size=10, bold=True)
    r2 = p.add_run(f'  —  {date}')
    set_font(r2, size=9, color=(100, 100, 100))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(school)
    set_font(r, size=9, italic=True, color=(80, 80, 80))


def add_cert(doc, name, sub, status):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(name)
    set_font(r1, size=10, bold=True)
    r2 = p.add_run(f'  —  {status}')
    set_font(r2, size=9, color=(100, 100, 100))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(sub)
    set_font(r, size=9, italic=True, color=(80, 80, 80))


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


# ─────────────────────────────────────────────
# ENGLISH CV
# ─────────────────────────────────────────────
def build_en():
    doc = Document()
    set_margins(doc)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('IKHLAS RETBI')
    set_font(r, name='Calibri', size=22, bold=True)

    # Title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run('Networks Engineer  ·  DevOps  ·  Cloud Security')
    set_font(r, size=10, color=(80, 80, 80))

    # Contacts
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)
    r = p3.add_run('ikhlasretbi@gmail.com  ·  +212 698 533 342  ·  Kenitra, Morocco  ·  linkedin.com/in/ikhlas-retbi  ·  github.com/ikhlas-rtb')
    set_font(r, size=9, color=(60, 60, 60))

    # Profile
    add_heading(doc, 'Profile')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        'Networks and Telecommunications Engineer with hands-on experience in DevOps automation, cloud infrastructure, '
        'and cybersecurity. Proven ability to design and deploy end-to-end CI/CD pipelines, containerized environments, '
        'and full monitoring stacks in production settings. Experienced with AWS, Terraform, Docker, Ansible, and the '
        'ELK Stack. Google Cybersecurity certified. Available immediately and open to relocation worldwide.'
    )
    set_font(r, size=9.5, color=(40, 40, 40))

    # Experience
    add_heading(doc, 'Experience')

    add_exp_item(doc,
        'DevOps & Automation Engineer — Intern (Final Year Thesis)',
        'Teligent SARL, Rabat',
        'Feb – Jun 2025',
        [
            'Deployed containerized services with Docker on the Teligent P90/E telecom platform in a POC environment',
            'Built end-to-end CI/CD pipelines with GitLab automating build, test, and validation cycles',
            'Implemented full monitoring stack: Grafana + Prometheus + ELK Stack for real-time supervision and centralized log analysis',
            'Automated infrastructure configuration with Ansible (IaC) ensuring reproducibility across environments',
            'Developed Python/Bash automation scripts with Makefile integration to optimize development workflows',
        ]
    )

    add_exp_item(doc,
        'IT Infrastructure & Systems — Intern',
        'Magna Mirrors Morocco, Kenitra',
        'Jul – Sep 2024',
        [
            'Administered Linux and Windows Server infrastructure with focus on availability and performance',
            'Monitored systems proactively via PRTG and performed log analysis for anomaly detection and incident resolution',
            'Managed user access via Active Directory and maintained full incident documentation',
        ]
    )

    add_exp_item(doc,
        'Avionics Systems Maintenance — Intern',
        'RAM Handling, Casablanca',
        'Apr – Jul 2023',
        [
            'Performed preventive maintenance on critical avionics systems (RADAR Boeing 737-NG) under strict protocols',
            'Produced technical intervention reports and maintained maintenance procedure documentation',
        ]
    )

    # Skills
    add_heading(doc, 'Technical Skills')
    add_skill_row(doc, 'DevOps & CI/CD', 'Docker, GitLab CI/CD, GitHub Actions, Ansible, Kubernetes, Git')
    add_skill_row(doc, 'Cloud & IaC', 'AWS (EC2, VPC, CloudTrail, CloudWatch, GuardDuty, Lambda), Terraform, Azure')
    add_skill_row(doc, 'Monitoring & Security', 'Grafana, Prometheus, ELK Stack, Splunk, PRTG, Azure Sentinel, SIEM, IDS/IPS')
    add_skill_row(doc, 'Scripting', 'Python, Bash, Shell, JavaScript')
    add_skill_row(doc, 'Systems & Networks', 'Linux (Ubuntu/CentOS/RHEL), Windows Server, Active Directory, TCP/IP, VLAN, GNS3, VMware, Proxmox')

    # Projects
    add_heading(doc, 'Projects')
    add_project(doc,
        'AWS Cloud Monitoring & Security Stack',
        'github.com/ikhlas-rtb/aws-cloud-monitoring',
        'CloudTrail + CloudWatch with 8 CIS Benchmark alarms + GuardDuty + Lambda auto-response + SNS alerting. Fully deployed and live-tested via CLI and Python.'
    )
    add_project(doc,
        'Terraform AWS Infrastructure Lab',
        'github.com/ikhlas-rtb/terraform-aws-lab',
        'Full AWS environment (VPC, EC2, Security Groups) via IaC with S3 remote state, modular structure, and GitHub Actions CI pipeline.'
    )
    add_project(doc,
        'SIEM Lab — Azure Sentinel',
        'github.com/ikhlas-rtb/SIEM-Lab-Project',
        'SIEM deployment on Microsoft Azure Sentinel with threat detection rules and centralized log analytics.'
    )
    add_project(doc,
        'DevOps Automation — Teligent P90/E (Thesis)',
        'github.com/ikhlas-rtb/P90E-project',
        'Full DevOps adoption on a telecom platform: Docker, GitLab CI/CD, Ansible IaC, Grafana + ELK monitoring stack.'
    )

    # Education
    add_heading(doc, 'Education')
    add_edu(doc,
        'MSc — Networks & Telecommunications Engineering',
        'Faculté des Sciences et Techniques, Settat  ·  Thesis: DevOps Adoption on Teligent P90/E via Ansible',
        '2023 – 2025'
    )
    add_edu(doc,
        'BSc — Electrical Engineering & Automated Systems',
        'Faculté des Sciences et Techniques, Settat',
        '2022 – 2023'
    )

    # Certifications
    add_heading(doc, 'Certifications')
    add_cert(doc,
        'Google Cybersecurity Professional Certificate',
        'Coursera / UM6P  ·  Network Security, SIEM, Linux Security, Incident Response',
        '2025'
    )
    add_cert(doc,
        'ALX Data Science & Data Engineering',
        'ALX Africa  ·  14-month program  ·  Started Feb 2025',
        'In Progress'
    )

    # Languages
    add_heading(doc, 'Languages')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Arabic (Native)  ·  French (Fluent)  ·  English (Professional)  ·  Spanish (Basic)')
    set_font(r, size=9.5)

    doc.save('Ikhlas_Retbi_CV_EN.docx')
    print('Saved: Ikhlas_Retbi_CV_EN.docx')


# ─────────────────────────────────────────────
# FRENCH CV
# ─────────────────────────────────────────────
def build_fr():
    doc = Document()
    set_margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('IKHLAS RETBI')
    set_font(r, name='Calibri', size=22, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run('Ingénieure Réseaux  ·  DevOps  ·  Sécurité Cloud')
    set_font(r, size=10, color=(80, 80, 80))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)
    r = p3.add_run('ikhlasretbi@gmail.com  ·  +212 698 533 342  ·  Kenitra, Maroc  ·  linkedin.com/in/ikhlas-retbi  ·  github.com/ikhlas-rtb')
    set_font(r, size=9, color=(60, 60, 60))

    add_heading(doc, 'Profil Professionnel')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        'Ingénieure en réseaux et télécommunications avec une expérience pratique en automatisation DevOps, '
        'infrastructure cloud et cybersécurité. Capacité avérée à concevoir et déployer des pipelines CI/CD complets, '
        'des environnements conteneurisés et des stacks de monitoring en conditions réelles. Maîtrise d\'AWS, Terraform, '
        'Docker, Ansible et de l\'ELK Stack. Certifiée Google Cybersecurity. Disponible immédiatement, mobilité mondiale.'
    )
    set_font(r, size=9.5, color=(40, 40, 40))

    add_heading(doc, 'Expérience Professionnelle')

    add_exp_item(doc,
        'Ingénieure DevOps & Automatisation — Stage PFE',
        'Teligent SARL, Rabat',
        'Fév – Juin 2025',
        [
            'Déploiement de services conteneurisés avec Docker sur la plateforme télécom Teligent P90/E en environnement POC',
            'Développement de pipelines CI/CD avec GitLab automatisant les cycles build, test et validation',
            'Implémentation d\'une stack de monitoring complète : Grafana + Prometheus + ELK Stack pour supervision temps réel et analyse centralisée de logs',
            'Automatisation des configurations via Ansible (IaC) garantissant la portabilité et la reproductibilité des environnements',
            'Développement de scripts Python/Bash avec intégration Makefile pour l\'optimisation des workflows de développement',
        ]
    )

    add_exp_item(doc,
        'IT Infrastructure & Systems — Stage',
        'Magna Mirrors Morocco, Kenitra',
        'Juil – Sept 2024',
        [
            'Administration d\'infrastructures Linux et Windows Server axée sur la disponibilité et la performance',
            'Supervision proactive via PRTG et analyse de logs système/réseau pour détection d\'anomalies et résolution d\'incidents',
            'Gestion des accès utilisateurs via Active Directory avec documentation complète des incidents',
        ]
    )

    add_exp_item(doc,
        'Maintenance Systèmes Avioniques — Stage',
        'RAM Handling, Casablanca',
        'Avr – Juil 2023',
        [
            'Maintenance préventive de systèmes avioniques critiques (RADAR Boeing 737-NG) selon protocoles stricts',
            'Rédaction de rapports techniques d\'intervention et documentation des procédures de maintenance',
        ]
    )

    add_heading(doc, 'Compétences Techniques')
    add_skill_row(doc, 'DevOps & CI/CD', 'Docker, GitLab CI/CD, GitHub Actions, Ansible, Kubernetes, Git')
    add_skill_row(doc, 'Cloud & IaC', 'AWS (EC2, VPC, CloudTrail, CloudWatch, GuardDuty, Lambda), Terraform, Azure')
    add_skill_row(doc, 'Monitoring & Sécurité', 'Grafana, Prometheus, ELK Stack, Splunk, PRTG, Azure Sentinel, SIEM, IDS/IPS')
    add_skill_row(doc, 'Scripting', 'Python, Bash, Shell, JavaScript')
    add_skill_row(doc, 'Systèmes & Réseaux', 'Linux (Ubuntu/CentOS/RHEL), Windows Server, Active Directory, TCP/IP, VLAN, GNS3, VMware, Proxmox')

    add_heading(doc, 'Projets')
    add_project(doc,
        'Stack Monitoring & Sécurité AWS',
        'github.com/ikhlas-rtb/aws-cloud-monitoring',
        'CloudTrail + CloudWatch (8 alarmes CIS) + GuardDuty + Lambda auto-réponse + alertes SNS. Déployé et testé en conditions réelles via CLI et Python.'
    )
    add_project(doc,
        'Infrastructure AWS avec Terraform',
        'github.com/ikhlas-rtb/terraform-aws-lab',
        'Environnement AWS complet (VPC, EC2, Groupes de sécurité) via IaC avec état distant S3, structure modulaire et pipeline CI GitHub Actions.'
    )
    add_project(doc,
        'SIEM Lab — Azure Sentinel',
        'github.com/ikhlas-rtb/SIEM-Lab-Project',
        'Déploiement SIEM sur Microsoft Azure Sentinel avec règles de détection de menaces et analyse centralisée de logs.'
    )
    add_project(doc,
        'Automatisation DevOps — Teligent P90/E (PFE)',
        'github.com/ikhlas-rtb/P90E-project',
        'Adoption DevOps complète sur plateforme télécom : Docker, GitLab CI/CD, Ansible IaC, stack monitoring Grafana + ELK.'
    )

    add_heading(doc, 'Formation')
    add_edu(doc,
        'Master Sciences et Techniques — Ingénierie des Réseaux',
        'FST Settat  ·  PFE : Adoption de la démarche DevOps sur la plateforme Teligent P90/E via Ansible',
        '2023 – 2025'
    )
    add_edu(doc,
        'Licence Sciences et Techniques — Génie Électrique',
        'FST Settat',
        '2022 – 2023'
    )

    add_heading(doc, 'Certifications')
    add_cert(doc,
        'Google Cybersecurity Professional Certificate',
        'Coursera / UM6P  ·  Sécurité réseau, SIEM, Linux, Réponse aux incidents',
        '2025'
    )
    add_cert(doc,
        'ALX Data Science & Data Engineering',
        'ALX Africa  ·  Programme 14 mois  ·  Débuté fév. 2025',
        'En cours'
    )

    add_heading(doc, 'Langues')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Arabe (Natif)  ·  Français (Courant)  ·  Anglais (Professionnel)  ·  Espagnol (Notions)')
    set_font(r, size=9.5)

    doc.save('Ikhlas_Retbi_CV_FR.docx')
    print('Saved: Ikhlas_Retbi_CV_FR.docx')


if __name__ == '__main__':
    build_en()
    build_fr()
    print('Done! Both CVs generated.')
